"""Generate ATS-friendly .docx resume — formatting matched to the user's existing template.

Format reverse-engineered from Data_Engineer_Deepika.docx:
- Page: A4, margins 0.75" top/bottom, 0.875" left/right
- Font: Times New Roman, 13pt body, 18pt name
- Colors: name + section headers in #1F3864, section underline in #2E5FA3
- Section headers ("Professional Summary:") have a thin navy bottom border
- Technical Skills is a 2-column table with light-gray (#CCCCCC) borders
- Bullets are ● with hanging indent
- Inline bold supported via **markdown-style** markers in summary/bullet text
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
BODY_PT = 13
NAME_PT = 18
ACCENT = RGBColor(0x1F, 0x38, 0x64)       # name + section header text
UNDERLINE_HEX = "2E5FA3"                  # section header bottom border
TABLE_BORDER_HEX = "CCCCCC"               # skills table borders


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def _set_run_font(run, size=BODY_PT, bold=False, italic=False, color=None, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), font)


def _set_para_spacing(p, before_twips=0, after_twips=0, line=None, line_rule="auto"):
    pPr = p._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(before_twips))
    spacing.set(qn("w:after"), str(after_twips))
    if line is not None:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), line_rule)


def _add_bottom_border(p, color_hex=UNDERLINE_HEX, size="8", space="2"):
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)


def _set_a4(doc):
    s = doc.sections[0]
    s.page_width = Twips(11906)
    s.page_height = Twips(16838)
    s.top_margin = Inches(0.75)
    s.bottom_margin = Inches(0.75)
    s.left_margin = Inches(0.875)
    s.right_margin = Inches(0.875)


def _set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(BODY_PT)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), FONT)


# ---------------------------------------------------------------------------
# inline bold parser
# ---------------------------------------------------------------------------

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_runs_with_inline_bold(p, text, *, base_bold=False, size=BODY_PT, color=None):
    """Append runs to paragraph p, bolding any **wrapped** segments. Other text
    inherits base_bold."""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            _set_run_font(r, size=size, bold=base_bold, color=color)
        r = p.add_run(m.group(1))
        _set_run_font(r, size=size, bold=True, color=color)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        _set_run_font(r, size=size, bold=base_bold, color=color)


# ---------------------------------------------------------------------------
# block builders
# ---------------------------------------------------------------------------

def _add_centered_line(doc, text, size=BODY_PT, bold=True, color=ACCENT, all_caps=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before_twips=0, after_twips=0)
    r = p.add_run(text.upper() if all_caps else text)
    _set_run_font(r, size=size, bold=bold, color=color)
    return p


def _add_section_header(doc, text):
    """13pt bold navy + bottom underline + 220 twips spacing before."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(p, before_twips=220, after_twips=40)
    r = p.add_run(f"{text}:")
    _set_run_font(r, size=BODY_PT, bold=True, color=ACCENT)
    _add_bottom_border(p)
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_para_spacing(p, before_twips=40, after_twips=0)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)

    bullet_run = p.add_run("●\t")
    _set_run_font(bullet_run, size=BODY_PT, bold=False)
    _add_runs_with_inline_bold(p, text, base_bold=False)

    p.paragraph_format.tab_stops.add_tab_stop(Inches(0.3))
    return p


def _header_block(doc, resume):
    _add_centered_line(doc, resume.get("name", "").upper(),
                       size=NAME_PT, bold=True, color=ACCENT)
    if resume.get("title"):
        _add_centered_line(doc, resume["title"],
                           size=BODY_PT, bold=True, color=ACCENT)

    contact = resume.get("contact", {})
    bits = []
    if contact.get("email"):
        bits.append(f"Email: {contact['email']}")
    if contact.get("phone"):
        bits.append(f"Mobile: {contact['phone']}")
    if contact.get("linkedin"):
        bits.append(contact["linkedin"])
    if contact.get("github"):
        bits.append(contact["github"])
    if bits:
        _add_centered_line(doc, "   |  ".join(bits),
                           size=BODY_PT, bold=True, color=ACCENT)


def _summary_block(doc, summary):
    if not summary:
        return
    _add_section_header(doc, "Professional Summary")
    for line in summary:
        _add_bullet(doc, line)


def _set_cell_borders(cell, color=TABLE_BORDER_HEX, size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), size)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        tcBorders.append(e)
    tcPr.append(tcBorders)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, v in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), str(v))
        e.set(qn("w:type"), "dxa")
        tcMar.append(e)
    tcPr.append(tcMar)


def _skills_block(doc, skills):
    """2-column table, light-gray borders, matching the original."""
    if not skills:
        return
    _add_section_header(doc, "Technical Skills")

    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    # column widths — sum should match content area width on A4 0.875" side margins:
    # 11906 - 2 * (0.875*1440) = 11906 - 2520 = 9386 ≈ 9360 dxa
    col_widths = [Inches(1.85), Inches(4.65)]   # ~2400 + 6960 dxa

    for category, items in skills.items():
        row = table.add_row()
        # category cell
        c0 = row.cells[0]
        c0.width = col_widths[0]
        _set_cell_borders(c0)
        _set_cell_margins(c0)
        p0 = c0.paragraphs[0]
        _set_para_spacing(p0, before_twips=20, after_twips=20)
        r0 = p0.add_run(category)
        _set_run_font(r0, size=BODY_PT, bold=True)

        # items cell
        c1 = row.cells[1]
        c1.width = col_widths[1]
        _set_cell_borders(c1)
        _set_cell_margins(c1)
        p1 = c1.paragraphs[0]
        _set_para_spacing(p1, before_twips=20, after_twips=20)
        items_text = ", ".join(items) if isinstance(items, list) else str(items)
        _add_runs_with_inline_bold(p1, items_text, base_bold=False)


def _experience_block(doc, experience):
    if not experience:
        return
    _add_section_header(doc, "Professional Experience")
    for role in experience:
        # company line
        company_loc = role.get("company", "")
        if role.get("location"):
            company_loc += f", {role['location']}"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_para_spacing(p, before_twips=120, after_twips=0)
        r = p.add_run(company_loc)
        _set_run_font(r, size=BODY_PT, bold=True)

        # title — dates
        title = role.get("title", "")
        dates = f"{role.get('start','')} – {role.get('end','')}".strip(" –")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_para_spacing(p, before_twips=0, after_twips=0)
        r = p.add_run(title)
        _set_run_font(r, size=BODY_PT, bold=True)
        if dates:
            r2 = p.add_run(f"\t({dates})")
            _set_run_font(r2, size=BODY_PT, bold=True)
            p.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT
            )

        for b in role.get("bullets", []):
            _add_bullet(doc, b)

        tech = role.get("tech")
        if tech:
            tech_str = ", ".join(tech) if isinstance(tech, list) else str(tech)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _set_para_spacing(p, before_twips=80, after_twips=0)
            r1 = p.add_run("Environment: ")
            _set_run_font(r1, size=BODY_PT, bold=True)
            r2 = p.add_run(tech_str)
            _set_run_font(r2, size=BODY_PT, bold=False)


def _education_block(doc, education):
    if not education:
        return
    _add_section_header(doc, "Education")
    for ed in education:
        line = ed.get("degree", "")
        if ed.get("school"):
            line += f" — {ed['school']}"
        if ed.get("year"):
            line += f" ({ed['year']})"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_para_spacing(p, before_twips=40, after_twips=0)
        r = p.add_run(line)
        _set_run_font(r, size=BODY_PT, bold=False)


def _certs_block(doc, certs):
    if not certs:
        return
    _add_section_header(doc, "Certifications")
    for c in certs:
        _add_bullet(doc, c)


def generate_docx(resume_data: dict, output_path: str | Path) -> Path:
    doc = Document()
    _set_default_style(doc)
    _set_a4(doc)

    _header_block(doc, resume_data)
    _summary_block(doc, resume_data.get("summary"))
    _skills_block(doc, resume_data.get("skills"))
    _experience_block(doc, resume_data.get("experience"))
    _education_block(doc, resume_data.get("education"))
    _certs_block(doc, resume_data.get("certifications"))

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
